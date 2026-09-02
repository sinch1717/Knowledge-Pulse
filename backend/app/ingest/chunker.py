"""Turning documents into chunks that remember where they came from.

The report's F3 exists because of a finding in the literature: splitting support
documents without regard for their structure measurably hurts retrieval. So the
splitter works in two stages. First it breaks content at heading boundaries, so a
chunk never straddles two unrelated sections. Then, only if a section is too long
for one chunk, it windows that section with an overlap.

Every chunk keeps its heading path — "Payments > Recurring invoices" — which is
what lets a citation name the section it came from instead of a page number.
"""

from __future__ import annotations

from dataclasses import dataclass, field

from bs4 import BeautifulSoup

from app.config import settings

HEADINGS = ("h1", "h2", "h3", "h4")
BLOCKS = ("p", "li", "pre", "td", "th", "blockquote", "dd", "dt")


@dataclass
class Section:
    heading_path: str
    text: str


@dataclass
class RawChunk:
    heading_path: str
    text: str
    url: str | None = None
    word_count: int = field(default=0)


# --------------------------------------------------------------------------
# Splitting into sections
# --------------------------------------------------------------------------

def sections_from_html(soup: BeautifulSoup) -> list[Section]:
    path: list[str] = []
    buffer: list[str] = []
    out: list[Section] = []

    def flush() -> None:
        body = " ".join(buffer).strip()
        if body:
            out.append(Section(heading_path=" › ".join(path) or "Introduction", text=body))
        buffer.clear()

    for element in soup.find_all(list(HEADINGS) + list(BLOCKS)):
        text = element.get_text(" ", strip=True)
        if not text:
            continue
        if element.name in HEADINGS:
            flush()
            level = int(element.name[1])
            path[:] = path[: level - 1]
            while len(path) < level - 1:
                path.append("")
            path.append(text)
            path[:] = [p for p in path if p]
        else:
            buffer.append(text)

    flush()
    return out


def sections_from_plain(text: str, title: str) -> list[Section]:
    """For PDFs and text files, where headings are unreliable, split on blank lines."""
    blocks = [b.strip() for b in text.split("\n\n") if b.strip()]
    return [Section(heading_path=title, text=b) for b in blocks]


# --------------------------------------------------------------------------
# Windowing long sections
# --------------------------------------------------------------------------

def chunk_sections(sections: list[Section], url: str | None = None) -> list[RawChunk]:
    target = settings.chunk_target_words
    overlap = settings.chunk_overlap_words
    chunks: list[RawChunk] = []

    # Merge consecutive short sections under the same heading rather than
    # emitting a chunk of nine words.
    merged: list[Section] = []
    for section in sections:
        if merged and merged[-1].heading_path == section.heading_path and len(
            merged[-1].text.split()
        ) < target // 2:
            merged[-1] = Section(section.heading_path, merged[-1].text + " " + section.text)
        else:
            merged.append(section)

    for section in merged:
        words = section.text.split()
        if len(words) <= target:
            if len(words) >= 15:  # a stub of a sentence is noise in the index
                chunks.append(
                    RawChunk(section.heading_path, section.text, url, len(words))
                )
            continue

        start = 0
        step = max(1, target - overlap)
        while start < len(words):
            window = words[start : start + target]
            if len(window) >= 15:
                chunks.append(
                    RawChunk(section.heading_path, " ".join(window), url, len(window))
                )
            start += step

    return chunks


# --------------------------------------------------------------------------
# File parsers
# --------------------------------------------------------------------------

def parse_pdf(path: str) -> str:
    import fitz  # pymupdf

    with fitz.open(path) as doc:
        return "\n\n".join(page.get_text() for page in doc)


def parse_docx(path: str) -> str:
    from docx import Document

    document = Document(path)
    parts: list[str] = []
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text.strip())
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n\n".join(parts)


def parse_text(path: str) -> str:
    with open(path, encoding="utf-8", errors="replace") as handle:
        return handle.read()
